from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs/project-level-documents")
FONT_CN = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    for line in text.strip().splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(64, 64, 64)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = FONT_CN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = FONT_CN
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    r.font.size = Pt(22)
    r.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.name = FONT_CN
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph("项目名称：汽车物流 WMS 仓库管理系统")
    doc.add_paragraph("适用对象：课程设计、项目答辩、开发交付、后续维护")
    doc.add_paragraph("生成日期：2026-06-29")
    doc.add_page_break()
    return doc


def build_requirements():
    doc = setup_doc(
        "汽车物流 WMS 仓库管理系统需求分析文档",
        "覆盖 PC 端、移动端、后端服务、数据库与权限体系",
    )

    doc.add_heading("一、项目背景与建设目标", level=1)
    doc.add_paragraph(
        "汽车物流仓储场景具有物料种类多、供应商多、批次多、扫码频繁、出入库追溯要求高等特点。"
        "传统人工台账方式难以及时反映库存状态，也难以保障 FIFO、封存、转包、盘点等现场操作的准确性。"
        "本系统通过 PC 管理端、移动扫码端和 Spring Boot 后端服务，构建覆盖基础资料、入库、出库、库存、转包、封存、盘点、预警和 AI 分析的 WMS 业务闭环。"
    )
    doc.add_paragraph(
        "建设目标包括：统一主数据、规范出入库流程、支持扫码作业、形成库存实时账、提供异常预警、保留审计追溯，并为答辩和后续扩展提供清晰的模块化代码结构。"
    )

    doc.add_heading("二、用户角色与使用场景", level=1)
    add_table(
        doc,
        ["角色", "典型用户", "核心诉求", "主要功能"],
        [
            ["admin", "系统管理员", "维护账号、角色、系统基础配置，查看全量数据", "用户权限管理、基础资料、审计日志、全部业务模块"],
            ["manager", "仓库经理/主管", "查看库存健康度、审核异常、掌握出入库情况", "仪表盘、库存报表、预警阈值、需求导入、审计日志"],
            ["operator", "仓库操作员", "快速完成扫码入库、出库、封存、转包、盘点", "移动端扫码、PC 日常业务操作、转包执行"],
        ],
    )

    doc.add_heading("三、系统范围", level=1)
    add_table(
        doc,
        ["端", "范围", "说明"],
        [
            ["PC Web", "管理和复核", "用于基础信息维护、入库单/出库单管理、库存报表、转包管理、封存管理、盘点任务、权限管理、AI 助手等。"],
            ["移动端", "现场扫码作业", "用于仓库现场扫码入库、扫码出库、带单出库、退库、封存/解封、转包、盘点等。"],
            ["后端服务", "业务规则和数据一致性", "提供 REST API、JWT 鉴权、事务控制、库存扣减、单据状态流转、审计日志和 AI 分析能力。"],
            ["数据库", "业务数据持久化", "存储用户、物料、供应商、包装、入库、出库、库存、看板、转包、封存、盘点、预警等数据。"],
        ],
    )

    doc.add_heading("四、功能性需求", level=1)
    add_table(
        doc,
        ["模块", "需求说明"],
        [
            ["登录与权限", "用户可通过账号密码登录；系统返回 JWT；后续请求携带 Token；管理员可创建用户、修改角色、启停账号；非管理员不能访问管理接口。"],
            ["基础信息管理", "维护物料、供应商、包装和库区等主数据，为入库、出库、库存统计提供基础。"],
            ["入库管理", "支持创建入库单、维护入库明细、扫码收货、生成看板标签、更新库存和入库历史。"],
            ["出库管理", "支持创建出库单、带单出库、扫码出库、FIFO 分配、库存扣减、出库历史追溯和退库。"],
            ["库存管理", "展示库存快照，支持按物料、供应商、库区查询，关联高低储阈值进行预警。"],
            ["需求管理", "支持需求批次和需求明细管理，为计划性出库提供依据。"],
            ["高低储预警", "维护物料安全库存阈值，根据库存状态展示低储、高储、正常等风险状态。"],
            ["封存/解封", "对已入库看板执行封存或解封，封存后不可出库和转包，确保异常物料隔离。"],
            ["转包管理", "支持拆包与合包。拆包会创建目标看板和入库单；合包会将数量汇入已有目标看板；所有操作记录转包历史。"],
            ["盘点管理", "支持创建盘点任务、移动端扫码盘点、记录盘点明细、统计差异，为库存校准提供依据。"],
            ["操作审计", "通过 AOP 自动记录新增、修改、删除操作，支持按用户、动作、对象追溯。"],
            ["AI 仓库助手", "基于库存和出入库历史进行缺货预测、呆滞预警，并支持自然语言查询。"],
            ["报表导出", "支持库存、入库、出库等数据导出为 Excel，方便线下汇报和归档。"],
        ],
    )

    doc.add_heading("五、核心业务流程需求", level=1)
    flows = {
        "登录认证流程": "用户输入账号密码 -> 前端表单校验 -> POST /api/auth/login -> BCrypt 校验密码 -> 生成 JWT -> 前端保存 Token -> 后续请求自动携带 Authorization",
        "入库流程": "创建入库单 -> 维护明细 -> 扫码/手工收货 -> 生成/更新看板标签 -> 增加库存 -> 更新单据状态 -> 记录入库历史",
        "出库流程": "创建出库单 -> 选择物料和数量 -> FIFO 查询可用看板 -> 扫码确认 -> 扣减库存和看板数量 -> 更新单据状态 -> 记录出库历史",
        "转包流程": "选择源看板 -> 校验已入库/未封存/可用数量 -> 判断目标看板是否存在 -> 拆包或合包 -> 原子扣减源数量 -> 创建出库/入库单据 -> 记录转包历史",
        "盘点流程": "创建盘点任务 -> 生成盘点明细 -> 移动端扫码录入实盘数 -> 计算差异 -> 任务完成 -> 管理端复核差异",
    }
    for title, flow in flows.items():
        doc.add_heading(title, level=2)
        add_code_block(doc, flow)

    doc.add_heading("六、非功能性需求", level=1)
    add_table(
        doc,
        ["类别", "需求"],
        [
            ["安全性", "密码使用 BCrypt 加密；除登录外接口均需 Token；管理员接口后端强制鉴权；敏感配置应放入本地配置或环境变量。"],
            ["一致性", "库存扣减、看板数量变化、单据创建、历史记录需要放入事务中，避免部分成功。"],
            ["可追溯性", "关键业务操作必须记录历史，包含操作人、时间、源对象、目标对象和数量变化。"],
            ["易用性", "PC 端提供清晰菜单和表格筛选；移动端优先扫码和大按钮交互；转包入口必须显性。"],
            ["可维护性", "采用 Controller、Service、Repository、DTO、Entity 分层结构；模块按业务目录组织。"],
            ["可运行性", "后端固定使用 JDK 21；前端和移动端通过 npm 脚本构建；数据库初始化脚本可重复执行。"],
        ],
    )

    doc.add_heading("七、数据需求", level=1)
    add_table(
        doc,
        ["数据对象", "核心字段", "用途"],
        [
            ["sys_user", "username、password、role、status", "登录认证、权限控制、用户管理"],
            ["material_info", "material_no、material_name、supplier_code", "物料主数据"],
            ["supplier_info", "supplier_code、supplier_name、contact、phone", "供应商主数据"],
            ["packaging_info", "material_no、package_model、package_capacity", "包装和器具容量"],
            ["inbound_order / detail", "doc_no、supplier、planned_qty、actual_qty、status", "入库单据和明细"],
            ["outbound_order / detail", "doc_no、supplier、planned_qty、actual_qty、status", "出库单据和明细"],
            ["inventory_stock", "material_code、supplier、warehouse_area、on_hand_qty", "库存快照"],
            ["inbound_kanban_label", "kanban_no、label_qty、sealed、transfer_status", "看板标签、扫码作业和转包封存控制"],
            ["package_transfer", "source_kanban_no、target_kanban_no、transfer_qty、transfer_type", "转包历史追溯"],
            ["inventory_check_task / detail", "task_no、system_qty、actual_qty、diff_qty", "盘点任务和差异"],
            ["audit_log", "username、action、target、detail、created_at", "操作审计"],
            ["ai_alert", "alert_type、risk_level、suggestion", "AI 预警结果"],
        ],
    )

    doc.add_heading("八、验收标准", level=1)
    add_numbers(
        doc,
        [
            "管理员、经理、操作员可按角色登录，非授权用户无法访问管理员接口。",
            "入库、出库、库存、封存、转包、盘点等核心流程可完成闭环。",
            "PC 前端、移动端 H5 和后端测试均可成功构建或运行。",
            "关键操作具备历史记录或审计记录，能够用于答辩演示和问题追溯。",
            "项目文档能说明需求、设计、测试和核心业务逻辑。",
        ],
    )
    doc.save(OUT_DIR / "WMS-requirements-analysis.docx")


def build_design():
    doc = setup_doc(
        "汽车物流 WMS 仓库管理系统设计文档",
        "系统架构、模块设计、数据库设计、接口设计与关键逻辑",
    )
    doc.add_heading("一、总体架构设计", level=1)
    doc.add_paragraph(
        "系统采用前后端分离架构，PC Web 与移动端通过 HTTP 调用 Spring Boot REST API。"
        "后端通过 Spring Security + JWT 做无状态认证，通过 Spring Data JPA 访问 MySQL。"
        "前端使用 Vue 3、Element Plus、Pinia、Vue Router；移动端使用 uni-app、Vue 3、Vant 风格组件。"
    )
    add_code_block(
        doc,
        """
PC Web(Vue3)        移动端(uni-app)
       \\              /
        \\ HTTP + JWT /
         Spring Boot REST API
          | Controller
          | Service 事务与业务规则
          | Repository JPA
          v
        MySQL 数据库
        """,
    )

    doc.add_heading("二、分层设计", level=1)
    add_table(
        doc,
        ["层次", "职责", "典型文件"],
        [
            ["Controller", "接收 HTTP 请求，参数校验，返回统一 ApiResult", "AuthController、TransferController、AdminUserController"],
            ["Service", "承载业务规则、事务控制、状态流转", "TransferServiceImpl、InboundOrderServiceImpl、OutboundOrderServiceImpl"],
            ["Repository", "封装数据库访问和自定义查询", "InboundKanbanLabelRepository、InventoryStockRepository"],
            ["Entity", "映射数据库表结构", "InboundKanbanLabel、PackageTransfer、SysUser"],
            ["DTO", "前后端传输对象，隔离实体和接口", "TransferResultDTO、ManagedUserDTO、InboundOrderSummaryDTO"],
            ["Frontend API", "封装 Axios 请求", "frontend/src/api/*.js、mobile/src/api/*.js"],
            ["View/Page", "承载用户交互和状态展示", "Transfer.vue、UserPermission.vue、mobile transfer/index.vue"],
        ],
    )

    doc.add_heading("三、认证与权限设计", level=1)
    doc.add_paragraph(
        "登录成功后后端生成 JWT，Token 中包含用户名和角色。前端将 Token 存入 localStorage，Axios 请求拦截器自动添加 Authorization 请求头。后端 JwtAuthFilter 解析 Token 并注入 SecurityContext。"
    )
    add_code_block(doc, "Authorization: Bearer <token>\nJwtAuthFilter -> jwtUtil.getUsernameFromToken() / getRoleFromToken() -> ROLE_admin 等权限对象")
    add_table(
        doc,
        ["权限点", "前端控制", "后端控制"],
        [
            ["登录", "白名单 /login", "/api/auth/login permitAll"],
            ["用户权限管理", "路由 meta.roles = admin，侧边栏仅 admin 显示", "/api/admin/** hasRole(\"admin\")"],
            ["基础资料和需求导入", "部分菜单 admin/manager 可见", "当前以认证为主，可继续扩展方法级权限"],
            ["日常业务操作", "登录后可访问", "需要有效 JWT"],
        ],
    )

    doc.add_heading("四、主要模块设计", level=1)
    add_table(
        doc,
        ["模块", "PC 端设计", "移动端设计", "后端设计"],
        [
            ["基础资料", "物料、供应商、包装、库区维护页面", "暂无", "BasicDataController + BasicDataService"],
            ["入库", "入库单管理、手工入库、入库历史", "扫码入库", "InboundOrderController、InboundScanController"],
            ["出库", "出库单管理、手工出库、出库历史、退库扫码", "扫码出库、带单出库、退库", "OutboundOrderController、OutboundScanController、OutboundReturnController"],
            ["库存", "库存报表、需求导入", "暂无", "InventoryController、DemandController"],
            ["封存", "封存/解封管理", "封存/解封扫码", "SealController、SealServiceImpl"],
            ["转包", "新建转包、可转包看板列表、转包历史", "源看板扫码、数量输入、确认转包", "TransferController、TransferServiceImpl"],
            ["盘点", "盘点任务管理", "盘点扫码和明细录入", "CheckController、CheckServiceImpl"],
            ["权限", "用户权限管理", "登录状态保存", "AdminUserController、UserServiceImpl、SecurityConfig"],
            ["AI", "AI 对话页、仪表盘预警卡片", "暂无", "AiChatController、AiScheduler、AiAlertService"],
        ],
    )

    doc.add_heading("五、转包详细设计", level=1)
    doc.add_paragraph("转包是本系统较复杂的库存调整功能，分为拆包和合包两类。它既要改变看板数量，又要同步库存和单据，同时记录历史。")
    add_table(
        doc,
        ["步骤", "设计要点"],
        [
            ["1. 查询源看板", "使用 findByKanbanNoWithLock 加悲观锁，防止并发操作同一源看板。"],
            ["2. 校验源状态", "源看板必须已入库、未封存、未全部转包，转包数量必须大于 0 且不超过可用数量。"],
            ["3. 判断模式", "目标看板号存在则合包，不存在或为空则拆包。"],
            ["4. 原子扣减", "decreaseLabelQty 在数据库层扣减源看板数量，失败则提示重试。"],
            ["5. 单据联动", "拆包创建出库单、入库单、目标看板；合包创建出库单并增加目标看板数量。"],
            ["6. 库存联动", "先扣减源库存，再按目标处理补回库存，保持物料总量一致。"],
            ["7. 历史记录", "保存 PackageTransfer，包含源/目标看板、数量、单据号、转包类型。"],
        ],
    )

    doc.add_heading("六、数据库设计", level=1)
    add_table(
        doc,
        ["表", "设计说明", "关键约束"],
        [
            ["sys_user", "用户账号、角色、状态", "username 唯一，password BCrypt 加密"],
            ["inbound_order / inbound_order_detail", "入库单头和明细", "doc_no 唯一，明细关联单头"],
            ["outbound_order / outbound_order_detail", "出库单头和明细", "doc_no 唯一，明细关联单头"],
            ["inventory_stock", "库存快照", "material_code + supplier + warehouse_area 唯一"],
            ["inbound_kanban_label", "看板标签和二维码载体", "kanban_no 唯一，记录 sealed、frozen_qty、transfer_status"],
            ["package_transfer", "转包历史", "source_kanban_no、target_kanban_no 建索引"],
            ["inventory_check_task / detail", "盘点任务与明细", "任务编号唯一，明细记录账面数和实盘数"],
            ["audit_log", "操作日志", "按用户、动作、时间查询"],
            ["ai_alert", "AI 预警", "按预警类型和风险等级查询"],
        ],
    )

    doc.add_heading("七、接口设计原则", level=1)
    add_bullets(
        doc,
        [
            "所有接口统一返回 ApiResult，包含 code、message、data、timestamp。",
            "业务异常通过 GlobalExceptionHandler 统一转换为 JSON 响应。",
            "分页接口返回 total、page、size、records，便于前端表格和移动端列表复用。",
            "新增、修改、删除类接口应记录操作日志，便于审计追溯。",
        ],
    )
    add_table(
        doc,
        ["接口组", "示例路径", "用途"],
        [
            ["认证", "/api/auth/login、/api/auth/userInfo", "登录、获取当前用户信息、修改密码"],
            ["权限", "/api/admin/users", "管理员维护用户和角色"],
            ["入库", "/api/inbound/**", "入库单、扫码入库、入库历史"],
            ["出库", "/api/outbound/**", "出库单、扫码出库、退库"],
            ["转包", "/api/transfer/kanbans、/api/transfer/execute", "可转包看板、执行转包、历史查询"],
            ["盘点", "/api/check/**", "盘点任务、盘点明细、扫码盘点"],
            ["AI", "/api/ai/chat、/api/ai/alerts/latest", "AI 对话与预警数据"],
        ],
    )

    doc.add_heading("八、运行环境设计", level=1)
    add_table(
        doc,
        ["组件", "版本/选择", "说明"],
        [
            ["JDK", "21", "通过 backend/mvn-jdk21.ps1 固定使用 Microsoft JDK 21.0.8"],
            ["Spring Boot", "2.7.18", "当前项目基础框架"],
            ["Node.js", "本机 npm 可用", "PC 前端和移动端依赖安装与构建"],
            ["MySQL", "5.7+ / 8.x", "schema.sql 初始化表结构和种子数据"],
            ["构建命令", "mvn test / npm run build / npm run build:h5", "分别验证后端、PC 前端、移动端"],
        ],
    )
    doc.save(OUT_DIR / "WMS-system-design.docx")


def build_testing():
    doc = setup_doc(
        "汽车物流 WMS 仓库管理系统测试文档",
        "测试范围、测试策略、测试用例、执行结果与风险说明",
    )
    doc.add_heading("一、测试目标", level=1)
    doc.add_paragraph(
        "本测试文档用于验证汽车物流 WMS 仓库管理系统是否满足需求分析和设计文档中的功能、权限、流程和运行环境要求。测试覆盖后端服务、PC 前端、移动端 H5、数据库初始化、核心业务流程和异常场景。"
    )

    doc.add_heading("二、测试范围", level=1)
    add_table(
        doc,
        ["范围", "测试内容"],
        [
            ["后端", "编译、单元测试、集成测试、权限接口、转包事务、出入库流程、盘点流程"],
            ["PC 前端", "登录、菜单权限、转包入口、用户权限管理、页面构建"],
            ["移动端", "扫码页面构建、转包页面、封存/解封、入库/出库页面"],
            ["数据库", "schema.sql 可执行，关键表字段和索引满足业务需求"],
            ["文档", "需求分析、设计、测试文档是否能支撑学习和答辩"],
        ],
    )

    doc.add_heading("三、测试环境", level=1)
    add_table(
        doc,
        ["项目", "环境"],
        [
            ["操作系统", "Windows 11"],
            ["后端 JDK", "Microsoft JDK 21.0.8，路径 C:\\Users\\86139\\.jdks\\ms-21.0.8"],
            ["Maven", "3.9.10"],
            ["PC 前端", "Vue 3 + Vite + Element Plus"],
            ["移动端", "uni-app + Vue 3"],
            ["数据库", "MySQL，默认库名 wms_db"],
        ],
    )

    doc.add_heading("四、自动化/构建测试结果", level=1)
    add_table(
        doc,
        ["编号", "命令", "预期结果", "实际结果", "结论"],
        [
            ["T-BE-001", "cd backend; .\\mvn-jdk21.ps1 test", "后端编译通过，测试通过", "Tests run: 76, Failures: 0, Errors: 0, Skipped: 2；BUILD SUCCESS", "通过"],
            ["T-BE-002", "cd backend; mvn test（默认 Java 24）", "如环境兼容则通过", "Lombok/JDK24 触发 TypeTag UNKNOWN 编译错误", "不通过，要求使用 JDK21"],
            ["T-FE-001", "cd frontend; npm install", "依赖安装完成", "added 121 packages", "通过"],
            ["T-FE-002", "cd frontend; npm run build", "PC 前端构建成功", "vite build 成功，生成 dist", "通过"],
            ["T-MO-001", "cd mobile; npm run build:h5", "移动端 H5 构建成功", "DONE Build complete", "通过"],
            ["T-GIT-001", "git push origin dev", "提交推送到 GitHub", "网络连接 GitHub 443 失败，Connection reset / Could not connect", "本地提交成功，远程推送待网络恢复"],
        ],
    )

    doc.add_heading("五、功能测试用例", level=1)
    add_table(
        doc,
        ["编号", "用例", "操作步骤", "预期结果", "当前状态"],
        [
            ["F-001", "登录成功", "输入 admin/admin123", "进入仪表盘，保存 Token 和用户信息", "待手工复测"],
            ["F-002", "登录失败", "输入错误密码", "提示用户名或密码错误，不进入系统", "待手工复测"],
            ["F-003", "管理员用户管理", "admin 进入用户权限页面，新增 operator", "新增成功，列表出现新用户", "待手工复测"],
            ["F-004", "角色修改", "将 operator 改为 manager", "角色更新成功，重新登录后权限变化", "待手工复测"],
            ["F-005", "账号禁用", "禁用某用户后用该用户登录", "登录失败，提示账号禁用", "待手工复测"],
            ["F-006", "非管理员访问权限接口", "manager/operator 调用 /api/admin/users", "返回 403", "待手工复测"],
            ["F-007", "入库单创建", "选择供应商和物料，输入计划数量", "生成入库单和明细，状态为未入库", "待手工复测"],
            ["F-008", "扫码入库", "扫描入库看板并确认数量", "库存增加，入库单状态更新", "待手工复测"],
            ["F-009", "出库单创建", "选择物料并输入计划出库数量", "生成出库单和明细", "待手工复测"],
            ["F-010", "扫码出库", "扫描可用看板进行出库", "库存扣减，FIFO 分配记录生成", "待手工复测"],
            ["F-011", "封存看板", "选择已入库看板执行封存", "看板 sealed=true，后续不可出库/转包", "待手工复测"],
            ["F-012", "解封看板", "对封存看板执行解封", "看板恢复可操作状态", "待手工复测"],
            ["F-013", "转包拆包", "源看板已入库未封存，目标看板号留空", "源数量减少，创建新目标看板、出库单、入库单、转包历史", "后端测试覆盖，待页面复测"],
            ["F-014", "转包合包", "填写已存在同物料目标看板号", "源数量减少，目标数量增加，记录转包历史", "后端测试覆盖，待页面复测"],
            ["F-015", "转包异常", "源看板封存、数量超限、目标物料不一致", "接口拒绝并返回明确错误", "后端测试覆盖"],
            ["F-016", "盘点任务", "创建任务并扫码录入实盘数", "生成盘点明细，计算差异", "后端测试覆盖，待页面复测"],
            ["F-017", "库存预警", "配置低储/高储阈值", "库存报表显示对应风险状态", "待手工复测"],
            ["F-018", "审计日志", "执行新增/修改/删除操作", "audit_log 记录用户、动作、对象和详情", "后端测试覆盖部分场景"],
            ["F-019", "AI 预警", "访问 AI 预警或对话接口", "返回库存分析或提示 API Key 配置问题", "待配置真实 key 后复测"],
            ["F-020", "移动端转包", "H5 打开转包页面，扫描/输入源看板", "展示源看板和转包表单，执行后显示结果", "构建通过，待真机/浏览器复测"],
        ],
    )

    doc.add_heading("六、接口测试重点", level=1)
    add_table(
        doc,
        ["接口", "测试点"],
        [
            ["POST /api/auth/login", "正确账号、错误密码、禁用账号、返回 Token 与 userInfo"],
            ["GET /api/auth/userInfo", "携带 Token 返回当前用户；无 Token 返回 401"],
            ["GET /api/admin/users", "admin 成功；manager/operator 返回 403"],
            ["POST /api/admin/users", "用户名唯一、密码加密、角色白名单、默认启用"],
            ["PUT /api/admin/users/{id}/role", "只允许 admin/manager/operator，非法角色报错"],
            ["PUT /api/admin/users/{id}/status", "只允许 0/1，禁用后不能登录"],
            ["GET /api/transfer/kanbans", "分页、物料筛选、供应商筛选，只返回可转包看板"],
            ["POST /api/transfer/execute", "拆包、合包、并发扣减、异常校验、历史记录"],
            ["GET /api/transfer/history", "分页、源看板筛选、目标看板筛选"],
        ],
    )

    doc.add_heading("七、缺陷与风险记录", level=1)
    add_table(
        doc,
        ["风险", "影响", "处理建议"],
        [
            ["默认 Java 24 无法编译后端", "开发者直接 mvn test 会失败", "使用 backend/mvn-jdk21.ps1，或全局 JAVA_HOME 切换到 JDK 21"],
            ["GitHub 推送失败", "本地 commit 暂未同步远端", "网络恢复后重新执行 git push origin dev"],
            ["手机端未真机扫码验证", "扫码体验可能与 H5 构建结果不同", "使用真实扫码设备或手机浏览器进行现场复测"],
            ["AI API Key 为 placeholder", "AI 对话可能无法真实调用 DeepSeek", "在 application-local.yml 配置真实 Key，并避免提交到 Git"],
            ["权限粒度仍是角色级", "无法精确到按钮级或字段级", "后续可设计 permission 表和 role_permission 表"],
        ],
    )

    doc.add_heading("八、测试结论", level=1)
    doc.add_paragraph(
        "当前项目在 JDK 21 环境下后端测试通过，PC 前端和移动端均可构建。核心业务模块具备演示和答辩基础，转包、权限管理和 JDK 运行问题已经完成关键修复。后续重点是进行完整手工联调、真机扫码测试、远程推送恢复以及权限粒度细化。"
    )
    doc.save(OUT_DIR / "WMS-test-document.docx")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_requirements()
    build_design()
    build_testing()
    for file in sorted(OUT_DIR.glob("*.docx")):
        print(f"{file} {file.stat().st_size}")


if __name__ == "__main__":
    main()

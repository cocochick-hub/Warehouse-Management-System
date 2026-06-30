# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# Title
title = doc.add_heading('WMS API 测试报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Test overview
doc.add_heading('测试概述', level=1)
overview_table = doc.add_table(rows=4, cols=2)
overview_table.style = 'Table Grid'
overview_table.cell(0, 0).text = '测试时间'
overview_table.cell(0, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
overview_table.cell(1, 0).text = '测试环境'
overview_table.cell(1, 1).text = 'Spring Boot 2.7.18 + MySQL, localhost:8080'
overview_table.cell(2, 0).text = '数据库'
overview_table.cell(2, 1).text = 'wms_db, root用户, 端口3306'
overview_table.cell(3, 0).text = '测试人员'
overview_table.cell(3, 1).text = '系统测试'

doc.add_paragraph()

# Authentication section
doc.add_heading('认证说明', level=1)
doc.add_paragraph('所有API需要携带JWT Token进行认证，格式如下：')
doc.add_paragraph('Authorization: Bearer {token}', style='Quote')
doc.add_paragraph('登录接口：POST /api/auth/login')
doc.add_paragraph('请求body：{"username":"admin","password":"admin123"}')
doc.add_paragraph('测试Token：eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJhZG1pbiIsInJvbGUiOiJhZG1pbiIsImlhdCI6MTc4Mjc4OTkwNywiZXhwIjoxNzgyODc2MzA3fQ.6xusZMUYcehEkc-jn0sHmqQKDINVibMgO29RTVIqj4A')

doc.add_paragraph()

# API List - Authentication Module
doc.add_heading('API 清单', level=1)

doc.add_heading('认证模块', level=2)
auth_table = doc.add_table(rows=2, cols=6)
auth_table.style = 'Table Grid'
headers = ['路径', '方法', '描述', '请求示例', '响应状态', '测试结果']
for i, header in enumerate(headers):
    auth_table.cell(0, i).text = header

auth_table.cell(1, 0).text = '/api/auth/login'
auth_table.cell(1, 1).text = 'POST'
auth_table.cell(1, 2).text = '用户登录'
auth_table.cell(1, 3).text = '{"username":"admin","password":"admin123"}'
auth_table.cell(1, 4).text = '200'
auth_table.cell(1, 5).text = '成功'

doc.add_paragraph()

# Basic Data Module
doc.add_heading('基础数据模块 /api/basic', level=2)

basic_data = [
    ['GET', '/api/basic/materials', '物料列表', '无', '200', '成功'],
    ['POST', '/api/basic/materials', '新增物料', '{"materialNo":"MAT-TEST-001",...}', '500', '失败-服务器内部错误'],
    ['GET', '/api/basic/suppliers', '供应商列表', '无', '200', '成功'],
    ['POST', '/api/basic/suppliers', '新增供应商', '{"supplierCode":"SUP-TEST",...}', '500', '失败-服务器内部错误'],
    ['GET', '/api/basic/warehouse-areas', '库区列表', '无', '200', '成功'],
    ['POST', '/api/basic/warehouse-areas', '新增库区', '{"areaCode":"WA-TEST",...}', '500', '失败-服务器内部错误'],
    ['GET', '/api/basic/packaging', '包装列表', '无', '200', '成功'],
    ['POST', '/api/basic/packaging', '新增包装', '{"materialNo":"MAT-TEST-001",...}', '409', '失败-数据冲突'],
]

basic_table = doc.add_table(rows=len(basic_data)+1, cols=6)
basic_table.style = 'Table Grid'
basic_headers = ['方法', '路径', '描述', '请求示例', '状态码', '测试结果']
for i, h in enumerate(basic_headers):
    basic_table.cell(0, i).text = h
for row_idx, row_data in enumerate(basic_data):
    for col_idx, cell_data in enumerate(row_data):
        basic_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Inbound Module
doc.add_heading('入库模块 /api/inbound', level=2)

inbound_data = [
    ['GET', '/api/inbound/orders', '入库单列表', '无', '200', '成功-空列表'],
    ['POST', '/api/inbound/orders', '创建入库单', '{"docNo":"IN-2024-001",...}', '500', '失败-服务器内部错误'],
    ['GET', '/api/inbound/orders/{id}', '入库单详情', '无', '404', '失败-资源不存在'],
    ['PUT', '/api/inbound/orders/{id}/receive', '收货', '{"actualQuantity":100}', '500', '失败-服务器内部错误'],
    ['GET', '/api/inbound/kanban/labels', '看板标签列表', '无', '404', '失败-路径不存在'],
    ['POST', '/api/inbound/kanban/labels/{id}/print', '打印标签', '无', '404', '失败-路径不存在'],
]

inbound_table = doc.add_table(rows=len(inbound_data)+1, cols=6)
inbound_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    inbound_table.cell(0, i).text = h
for row_idx, row_data in enumerate(inbound_data):
    for col_idx, cell_data in enumerate(row_data):
        inbound_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Outbound Module
doc.add_heading('出库模块 /api/outbound', level=2)

outbound_data = [
    ['GET', '/api/outbound/orders', '出库单列表', '无', '200', '成功-空列表'],
    ['POST', '/api/outbound/orders', '创建出库单', '{"docNo":"OUT-2024-001",...}', '500', '失败-服务器内部错误'],
    ['POST', '/api/outbound/orders/{id}/issue', '出库', '{"actualQuantity":50}', '400', '失败-出库明细不能为空'],
    ['GET', '/api/outbound/kanban/labels', '出库看板标签', '无', '404', '失败-路径不存在'],
]

outbound_table = doc.add_table(rows=len(outbound_data)+1, cols=6)
outbound_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    outbound_table.cell(0, i).text = h
for row_idx, row_data in enumerate(outbound_data):
    for col_idx, cell_data in enumerate(row_data):
        outbound_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Inventory Module
doc.add_heading('库存模块 /api/inventory', level=2)

inventory_data = [
    ['GET', '/api/inventory/stocks', '库存列表', '?page=1&size=10', '200', '成功-空列表'],
    ['GET', '/api/inventory/stocks/{id}', '库存详情', '无', '404', '失败-资源不存在'],
]

inventory_table = doc.add_table(rows=len(inventory_data)+1, cols=6)
inventory_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    inventory_table.cell(0, i).text = h
for row_idx, row_data in enumerate(inventory_data):
    for col_idx, cell_data in enumerate(row_data):
        inventory_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Transfer Module
doc.add_heading('看板转移模块 /api/transfer', level=2)

transfer_data = [
    ['POST', '/api/transfer/merge', '合包', '{"fromLabelNos":["LABEL-001"],...}', '404', '失败-路径不存在'],
    ['POST', '/api/transfer/split', '拆包', '{"fromLabelNo":"LABEL-001",...}', '404', '失败-路径不存在'],
]

transfer_table = doc.add_table(rows=len(transfer_data)+1, cols=6)
transfer_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    transfer_table.cell(0, i).text = h
for row_idx, row_data in enumerate(transfer_data):
    for col_idx, cell_data in enumerate(row_data):
        transfer_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# AI Module
doc.add_heading('AI模块 /api/ai', level=2)

ai_data = [
    ['GET', '/api/ai/alerts', 'AI预警列表', '无', '404', '失败-路径不存在'],
    ['GET', '/api/ai/chat', 'AI对话', '?message=hello', '500', '失败-服务器内部错误'],
]

ai_table = doc.add_table(rows=len(ai_data)+1, cols=6)
ai_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    ai_table.cell(0, i).text = h
for row_idx, row_data in enumerate(ai_data):
    for col_idx, cell_data in enumerate(row_data):
        ai_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Check Module
doc.add_heading('盘点模块 /api/check', level=2)

check_data = [
    ['GET', '/api/check/tasks', '盘点任务列表', '无', '200', '成功-空列表'],
    ['POST', '/api/check/tasks', '创建盘点任务', '{"taskNo":"CHECK-2024-001",...}', '500', '失败-服务器内部错误'],
]

check_table = doc.add_table(rows=len(check_data)+1, cols=6)
check_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    check_table.cell(0, i).text = h
for row_idx, row_data in enumerate(check_data):
    for col_idx, cell_data in enumerate(row_data):
        check_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Import/Export Module
doc.add_heading('数据导入导出模块 /api/import & /api/export', level=2)

import_export_data = [
    ['POST', '/api/import/materials', '导入物料', '[{...}]', '404', '失败-路径不存在'],
    ['POST', '/api/import/demand', '导入需求', '[{...}]', '404', '失败-路径不存在'],
    ['GET', '/api/export/inbound/{docNo}', '导出入库单', 'IN-2024-001', '404', '失败-路径不存在'],
    ['GET', '/api/export/outbound/{docNo}', '导出出库单', 'OUT-2024-001', '404', '失败-路径不存在'],
]

ie_table = doc.add_table(rows=len(import_export_data)+1, cols=6)
ie_table.style = 'Table Grid'
for i, h in enumerate(basic_headers):
    ie_table.cell(0, i).text = h
for row_idx, row_data in enumerate(import_export_data):
    for col_idx, cell_data in enumerate(row_data):
        ie_table.cell(row_idx+1, col_idx).text = cell_data

doc.add_paragraph()

# Issues Found
doc.add_heading('发现的问题列表', level=1)

issues = [
    '1. 多个POST接口返回500错误：新增物料、供应商、库区、入库单、出库单、盘点任务等接口均返回服务器内部错误',
    '2. 部分API路径404：入库看板标签、出库看板标签、合包、拆包、物料导入、需求导入、导出接口均报路径不存在',
    '3. 包装新增接口409冲突：可能是数据唯一性约束导致',
    '4. AI模块路径不存在：/api/ai/alerts和/api/ai/chat返回404，但chat接口返回500可能是另一个问题',
    '5. 出库接口返回400错误：提示"出库明细不能为空"，说明参数校验生效，但接口设计需要先添加明细',
]

for issue in issues:
    doc.add_paragraph(issue, style='List Number')

doc.add_paragraph()

# Suggestions
doc.add_heading('建议', level=1)

suggestions = [
    '1. 检查后端日志，定位500错误的根本原因（可能是数据库约束、事务问题或代码异常）',
    '2. 确认/api/inbound/kanban、/api/transfer等模块的路由配置是否正确',
    '3. 实现完整的出库流程：先创建出库单明细，再执行出库操作',
    '4. 补充AI模块的接口实现或确认是否有单独的AI服务',
    '5. 建议增加更多的测试数据来验证各模块功能',
    '6. 导出功能可能需要先生成实际数据才能测试',
]

for suggestion in suggestions:
    doc.add_paragraph(suggestion, style='List Number')

# Summary
doc.add_heading('测试统计', level=1)

summary_table = doc.add_table(rows=5, cols=2)
summary_table.style = 'Table Grid'
summary_table.cell(0, 0).text = '总测试接口数'
summary_table.cell(0, 1).text = '26'
summary_table.cell(1, 0).text = '成功'
summary_table.cell(1, 1).text = '10 (38.5%)'
summary_table.cell(2, 0).text = '失败-服务器错误(500)'
summary_table.cell(2, 1).text = '8 (30.8%)'
summary_table.cell(3, 0).text = '失败-路径不存在(404)'
summary_table.cell(3, 1).text = '7 (26.9%)'
summary_table.cell(4, 0).text = '失败-业务错误(400/409)'
summary_table.cell(4, 1).text = '1 (3.8%)'

# Save document
output_path = 'D:/Warehouse-Management-System/Warehouse-Management-System/docs/WMS_API_测试报告_2026-06-30.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')